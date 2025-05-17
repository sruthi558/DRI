# Importing flask module in the project is mandatory
# An object of Flask class is our WSGI application.
from flask import flash,render_template, url_for, redirect, session, send_file
import gridfs
from flask_mail import Mail, Message
from bson import ObjectId
from pymongo import MongoClient
from flask import Flask
from flask import request
import pandas as pd
from random import randint
from flask_paginate import Pagination
from datetime import timedelta,datetime
import os
import zipfile


cluster = MongoClient("mongodb://localhost:27017")
app = Flask(__name__)
app.config["UPLOAD_DIR"] = "uploads"
app.config["MAIL_SERVER"] = "smtp.gmail.com"
app.config["MAIL_USERNAME"] = "isra.government@gmail.com"
app.config["MAIL_PASSWORD"] = "zgdwdccbsxyrrvkx"
app.config["MAIL_PORT"] = 465
app.config["MAIL_USE_TLS"] = False
app.config["MAIL_USE_SSL"] = True
mail = Mail(app)


# Secret key for sessions encryption
app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'

# Connection to Database
conn = MongoClient("mongodb://localhost:27017")
db = conn.DRI

# The route() function of the Flask class is a decorator,
# which tells the application which URL should call
# the associated function.


@app.before_request
def make_session_permanent():
    session.permanent = True
    app.permanent_session_lifetime = timedelta(minutes=50)
    session.modified = True


@app.route("/")
def base():
    return redirect(url_for("login"))


@app.route("/registration")
def registration():
    session.clear()
    return render_template("admin_login.html")


@app.route("/login")
def login():
    message = None
    if "message" in session:
        message = session["message"]
    if "admin_msg" in session:
        message = session["admin_msg"]
    if "invalid_otp_msg" in session:
        message = session["invalid_otp_msg"]
    session.clear()
    dri_db = cluster["DRI"]
    if "formations" in list(dri_db.list_collection_names()):
        collection = dri_db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        formation_data_list = formation_data.values.tolist()
    else:
        formation_data_list=[]

    return render_template("login.html",collections_list = formation_data_list,msg=message)



@app.route("/post_registration", methods=["post", "get"])
def post_registration():
    if request.method == "POST":
        ssoid = request.form["ssoid"]
        password = request.form["password"]
        email = request.form["email"]
        designation = 'admin'
        verify_otp = request.form.get("verify_otp")
        # Earlier otp
        input_otp = request.form.get("otp")
        if str(input_otp) == str(verify_otp):
            db = cluster["Authentication"]
            collection = db["Admin"]
            collection.insert_one({"ssoid": ssoid, "password": password, "email": email,"designation":designation})
            session["admin_msg"] = "Profile Successfully Added"
            return redirect(url_for("login"))
        else:
            session["invalid_otp_msg"] = f"invalid OTP"
            return redirect(url_for("login"))



@app.route("/enter_otp", methods=["post", "get"])
def enter_otp():
    if request.method == "POST":
        ssoid = request.form.get("ssoid")
        password = request.form.get("password")
        designation = "admin"

        email = request.form.get("email")
        msg = Message(subject='OTP', sender='isra.government@gmail.com', recipients=[email])
        otp = randint(000000, 999999)
        msg.body = str(otp)
        mail.send(msg)
        return render_template("otp.html", credential={"email": str(email), "password": str(password), "ssoid": str(ssoid), "designation": str(designation),"verify_otp": otp})

        # Earlier code
        # return render_template("otp.html", credential={"email": str(email), "password": str(password), "ssoid": str(ssoid), 'designation':designation, "verify_otp": otp})


@app.route("/post_otp", methods=["post", "get"])
def post_otp():
    if request.method == "POST":
        ssoid = request.form.get("ssoid")
        password = request.form.get("password")
        email = request.form.get("email")
        designation = request.form.get("designation")
        verify_otp = request.form.get("verify_otp")
        # Earlier otp
        input_otp = request.form.get("otp")
        if str(input_otp) == str(verify_otp):
            collection = db["authentication"]
            document=collection.find_one( {"$and": [{"ssoid": ssoid}, {"password": password},{"email":email}]})
            if document["ssoid"] == ssoid and document["password"] == password and document["email"]==email:
                session["uid"] = str(document["_id"])
                session["designation"] = document["designation"]
                session["email_id"] = document["email"]
                session["unique_id"] = session["uid"] + ssoid
                print("session_databse:", session["uid"])
                # return redirect(url_for("home"))
                # Earlier code
                # return render_template("home.html", result = {"designation": session["designation"]})
                designation=session["designation"]
                if designation=="admin":
                    return render_template("home.html",designation=designation)
                else:
                    return render_template("home.html")
        else:
            msg = f"Please enter valid otp"
            return render_template("otp.html", msg=msg, credential={"email": str(email), "password": str(password), "ssoid": str(ssoid), "verify_otp": verify_otp})
        return redirect(url_for("login"))


# @app.route("/postsignin", methods=["post", "get"])
# def postsignin():
#     if request.method == "POST":
#         user_otp = request.form.get('otp')
#         collection = db["authentication"]
#         document = collection.update_one({"email":session["email"]}, { "$set": {"otp":str(user_otp)}})
#         try:
#             document = collection.find_one({"otp":str(otp)})
#             session["otp"] = str(document["otp"])
#             session["unique_id"] = session["uid"]+session["otp"]
#             document = collection.update_one({"email": session["email"]},{"$unset":{"otp":""}})
#         except Exception as e:
#             msg=f"Enter Correct OTP {e}"
#             return render_template("otp.html", msg=msg)
#         print(session["otp"])
#         return redirect(url_for("home"))


# @app.route("/postsignin", methods=["post", "get"])
# def postsignin():
#     try:
#         if request.method == "POST":
#             ssoid = request.form.get("ssoid")
#             password = request.form.get("password")
#             # designation = request.form.get("designation")

#             collection = db["authentication"]
#             document = collection.find_one(
#                 {"$and": [{"ssoid": ssoid}, {"password": password}]}
#             )
#             # print("username : ", document["ssoid"], "passw : ", document["password"])
#             if document["ssoid"] == ssoid and document["password"] == password:
#                 session["uid"] = str(document["_id"])
#                 session["designation"] = document["designation"]
#                 session["email_id"] = document["email"]
#                 session["unique_id"] = session["uid"] + ssoid
#                 print("session_database:", session["uid"])
#                 # return redirect(url_for("home"))
#                 if session["designation"] == "admin":
#                     return render_template("home.html", designation=session["designation"])
#             elif document["ssoid"] == ssoid and document["password"] == password:
#                 session["uid"] = str(document["_id"])
#                 session["designation"] = document["designation"]
#                 session["formation"] = document["formation"]

#                 session["email_id"] = document["email"]
#                 session["unique_id"] = session["uid"] + ssoid
#                 print("session_database:", session["uid"])
#                 # return redirect(url_for("home"))
#                 if session["formation"] == "zone_one":
#                     return render_template(".html", designation=session["designation"])
            
#                 # elif session["designation"] == "officer":
#                 #     return render_template("officer.html", designation=session["designation"])
#                 # elif session["designation"] == "zone1_officer":
#                 #     return render_template("zone_one_officer.html", designation=session["designation"])
#                 # else:
#                 #     return render_template("index.html", designation=session["designation"])

#             else:
#                 msg = "Invalid username or password"
#                 return render_template("login.html", msg=msg)
#     except Exception as e:
#         msg = f"Invalid username or password {e}"
#         return render_template("login.html", msg=msg)


@app.route("/postsignin", methods=["POST", "GET"])
def postsignin():
    try:
        if request.method == "POST":
            ssoid = request.form.get("ssoid")
            password = request.form.get("password")
            branch_level = str(request.form.get("branch_level"))
            db = cluster["Authentication"]
            collection = db[branch_level]
            document = collection.find_one({"ssoid": ssoid})
            if document["ssoid"] == ssoid and document["password"] == password:
                session["uid"] = str(document["_id"])
                session["designation"] = document["designation"]
                session["formation"] = str(branch_level)
                session["email_id"] = document["email"]
                # session["unique_id"] = session["uid"] + ssoid
                return redirect(url_for("home"))
            else:
                # message ='Invalid username or password'
                # url = url_for('login', message=message)
                # return redirect(url)
                session["message"] = 'Invalid username or password'
                return redirect(url_for("login"))

    except Exception as e:
        # message ='Invalid username or password'
        # url = url_for('login', message=message)
        # return redirect(url)
        session["message"] = 'Invalid username or password'
        return redirect(url_for("login"))
    

@app.route("/admin_postsignin", methods=["POST", "GET"])
def admin_postsignin():
    try:
        if request.method == "POST":
            ssoid = request.form.get("ssoid")
            password = request.form.get("password")
            db = cluster["Authentication"]
            collection = db["Admin"]
            document = collection.find_one({"ssoid": ssoid})
            if document["ssoid"] == ssoid and document["password"] == password:
                session["uid"] = str(document["_id"])
                session["designation"] = document["designation"]
                session["email_id"] = document["email"]
                session["formation"] = "Admin"
                # session["unique_id"] = session["uid"] + ssoid
                return redirect(url_for("home"))
            
            else:
                session["message"] = 'Invalid username or password'
                return redirect(url_for("login"))

    except Exception as e:
        session["message"] = 'Invalid username or password'
        return redirect(url_for("login"))

@app.route("/forgot-password")
def forgotpassword():
    return render_template("forgot-password.html")


@app.route("/post-forgot-password", methods=["post", "get"])
def postforgotpassword():
    email = request.form.get("email")
    try:
        collection = db["authentication"]
        document = collection.find_one({"$and": [{"email": email}]})
        msg = Message(
            "Account Credentials",
            sender="isra.government@gmail.com",
            recipients=[email],
        )
        msg.body = "Account password: {}".format(document["password"])
        mail.send(msg)
        msg = "Your credentials has been sent to your email address"
        return render_template("registration.html", msg=msg)

    except Exception as e:
        print(e)
        msg = f"Unregistered E-mail ID {e}"
        return render_template("forgot-password.html", msg=msg)


@app.route("/forgot-ssoid")
def forgotssoid():
    return render_template("forgot-ssoid.html")




@app.route("/post-forgot-ssoid", methods=["post", "get"])
def postforgotssoid():
    email = request.form.get("email")
    try:
        collection = db["authentication"]
        document = collection.find_one({"$and": [{"email": email}]})
        msg = Message(
            "Account Credentials",
            sender="isra.government@gmail.com",
            recipients=[email],
        )
        msg.body = "Account SSOID : {}".format(document["ssoid"])
        mail.send(msg)
        msg = "Your credentials has been sent to your email address"
        return render_template("registration.html", msg=msg)

    except Exception as e:
        print(e)
        msg = f"Unregistered E-mail ID {e}"
        return render_template("forgot-ssoid.html", msg=msg)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

def navbar_conditions():
    if session["designation"]=="officer":
        content = {"AU":False,"SI":False,"D1":True,"D2":True,"D3":True}
    elif session["designation"] == "Designated Deputy Director":
        content={"AU":False, "SI":True,"D1":False,"D2":True,"D3":True}
    elif session["designation"] == "Designated Additional Director":
        content = {"AU": False,"SI":True,"D1": False,"D2": True,"D3":True}
    elif session["designation"] == "DRI Zone-ADG":
        content = {"AU":False,"SI":True,"D1": True,"D2":True,"D3":True}
    elif session["designation"] == "Designated Deputy Director-HQRS":
        content = {"AU":False, "SI":True,"D1":False,"D2":True,"D3":True}
    elif session["designation"] == "Designated Additional Director-HQRS":
        content = {"AU":False, "SI":True,"D1": False,"D2": True,"D3":True}
    elif session["designation"] == "Pr.ADG-HQRS":
        content = {"AU": False, "SI":True,"D1":False,"D2": True,"D3":True}
    elif session["designation"] == "Director General":
        content = {"AU": False, "SI":True,"D1": True,"D2": True,"D3":True}
    elif session["designation"] == "admin":
        content = {"AU": True, "SI":False,"D1": False,"D2": False,"D3":False}
    return content

def button_conditions():
    if session["designation"]=="officer":
        buttons = {"WD1":True,"RD1":True,"WD2":True,"RD2":True,"RD3":True,"WD3":True}
    elif session["designation"] == "Designated Deputy Director":
        buttons = {"WS":True,"RS":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    elif session["designation"] == "Designated Additional Director":
        buttons = {"WS":True,"RS":True,"RD2":True,"WD2":False,"RD3":True,"WD3":False}
    elif session["designation"] == "DRI Zone-ADG":
        buttons = {"RS":True,"WS":False,"RD1":True,"WD1":False,"RD2":True,"WD2":False,"RD3":True,"WD3":False}
    elif session["designation"] == "Designated Deputy Director-HQRS":
        buttons = {"WS":False,"RS":True,"WD1":False,"RD1":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    elif session["designation"] == "Designated Additional Director-HQRS":
        buttons = {"WS":False,"RS":True,"WD1":False,"RD1":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    elif session["designation"] == "Pr.ADG-HQRS":
        buttons = {"WS":False,"RS":True,"WD1":False,"RD1":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    elif session["designation"] == "Director General":
        buttons = {"WS":False,"RS":True,"WD1":False,"RD1":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    elif session["designation"] == "admin":
        buttons = {"WS":False,"RS":True,"WD1":False,"RD1":True,"WD2":False,"RD2":True,"RD3":True,"WD3":False}
    return buttons




@app.route("/home")
def home():
    return render_template("home.html",content=navbar_conditions(), buttons=button_conditions())

@app.route("/DRI1")
def DRI1():
    return render_template("DRI1.html",content=navbar_conditions(),buttons=button_conditions())

@app.route("/DRI2")
def DRI2():
    return render_template("DRI2.html",content = navbar_conditions(),buttons=button_conditions())

@app.route("/DRI3")
def DRI3():
    return render_template("DRI3.html",content=navbar_conditions(),buttons=button_conditions())

@app.route("/shareintelligence")
def shareintelligence():
    return render_template("shareintelligence.html",content = navbar_conditions(),buttons=button_conditions())
    
@app.route("/user_page")
def user_page():
    return render_template("user_page.html",content = navbar_conditions(),buttons=button_conditions())

@app.route("/user_value",methods=["POST"])
def user_value():
    if request.method == "POST":
        button_req = request.form.get("button_req")
        if button_req == "Add User":
            return redirect(url_for("add_user"))
        elif button_req == "Edit User":
            return redirect(url_for("edit_user"))

@app.route("/create-new-DRI1")
def create_new_DRI1():
    print("session_id_create_new_DRI1:", session["unique_id"])
    return render_template("create_new_DRI1.html",buttons=button_conditions())


@app.route("/create-new-DRI2")
def create_new_DRI2():
    print("session_id_create_new_DRI2:", session["unique_id"])
    return render_template("create_new_DRI2.html",content=navbar_conditions(),buttons=button_conditions())

@app.route("/create_dri1",methods=["GET", "POST"])
def create_dri1():
    msg =None
    if "dri1_filled_msg" in session:
        msg = session.pop("dri1_filled_msg",None)
    if "error_msg" in session:
        msg = session.pop("error_msg",None)
    # if button_input == "Create New DRI1":
    collection = db["designation"]
    designation_data = pd.DataFrame(list(collection.find()))
    collection = db["commodity"]
    total_data = pd.DataFrame(list(collection.find()))
    collection = db["formations"]
    formation_data = pd.DataFrame(list(collection.find()))
    return render_template("create_new_DRI1.html",content=navbar_conditions(),AU=False,SI=False,designation=session["designation"],formation=session["formation"], filter_commodity=total_data.values.tolist(),filter_place = designation_data.values.tolist(),filter_formation=formation_data.values.tolist(),msg=msg)




@app.route("/create_dri2",methods=["GET", "POST"])
def create_dri2():
    button_input = request.form.get("submit_button")
    if button_input == "Create New DRI2":
        collection = db["designation"]
        designation_data = pd.DataFrame(list(collection.find()))
        collection = db["commodity"]
        total_data = pd.DataFrame(list(collection.find()))
        collection = db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        return render_template("create_new_DRI2.html",content=navbar_conditions(),SI=False,AU=False,designation=session["designation"],formation=session["formation"],  filter_commodity=total_data.values.tolist(),filter_place = designation_data.values.tolist(),filter_formation=formation_data.values.tolist())



def vi_col_dri1(col_name_list):
    db = cluster["DRI1"]
    total_data_final = pd.DataFrame()
    total = 0
    doc = {}
    for val in col_name_list:
        collection1 = db[str(val)]
        print("here")
        placedropdown = collection1.find({"Form": "DRI 1"}, {'place': True})
        # source_data = collection1.find({"Form": "Share inputs by your formation"}, {'id': False})
        # for i in source_data:
        #     print(i)
        dri1_allKeys = collection1.find({"Form": "DRI 1"}, {'_id': False, 'Form': False,'other_commodity':False, 'involve': False,'place':False, 'links': False, 'action': False, 'remarks': False, 'designation': False, 'DRI Zonal Unit': False, 'Upload Supporitng Documents': False })
        dri1_allKeys_df = pd.DataFrame(list(collection1.find({"Form": "DRI 1",}, {'_id': False, 'Form': False, 'other_commodity':False,'place':False,'involve': False, 'links': False, 'action': False, 'remarks': False, 'designation': False, 'DRI Zonal Unit': False, 'Upload Supporitng Documents': False })))
        total = total + len(dri1_allKeys_df)
        allKeys_df_dummy = pd.DataFrame(list(collection1.find({"Form": "DRI 1"}, {'Form': False, 'other_commodity':False,'involve': False, 'links': False,'action': False, 'remarks': False, 'designation': False, 'DRI Zonal Unit': False, 'Upload Supporitng Documents': False })))
        # allKeys_df_dummy["collection_id"] = id
        total_data_final = pd.concat([total_data_final, allKeys_df_dummy], ignore_index=True)
        for doc in dri1_allKeys:
            pass
        allKeys_df_dummy_list = total_data_final.to_dict(orient="records")
        return total,doc,allKeys_df_dummy_list,list(allKeys_df_dummy)
    

def vi_col_dri2(col_name_list):
    db = cluster["DRI2"]
    total_data_final = pd.DataFrame()
    total = 0
    doc = {}
    for val in col_name_list:
        collection2 = db[str(val)]
        print("here")
        # source_data = collection1.find({"Form": "Share inputs by your formation"}, {'id': False})
        # for i in source_data:
        #     print(i)
        dri1_allKeys = collection2.find({"Form": "DRI 2"}, {'_id':False, 'serialno': True, 'date': True,  'caseno': True ,'commissionarate': True, 'division': True })
        dri1_allKeys_df = pd.DataFrame(list(collection2.find({"Form": "DRI 2",}, {'_id':False, 'serialno': True, 'date': True, 'caseno': True ,'commissionarate': True, 'division': True })))
        total = total + len(dri1_allKeys_df)
        allKeys_df_dummy = pd.DataFrame(list(collection2.find({"Form": "DRI 2"}, {'_id':False, 'place':True,'serialno': True, 'date': True, 'caseno': True ,'commissionarate': True, 'division': True })))
        # allKeys_df_dummy["collection_id"] = id
        total_data_final = pd.concat([total_data_final, allKeys_df_dummy], ignore_index=True)
        for doc in dri1_allKeys:
            pass
        allKeys_df_dummy_list = total_data_final.to_dict(orient="records")
        return total,doc,allKeys_df_dummy_list,list(allKeys_df_dummy)

@app.route("/view-DRI1")
def view_DRI1():
    allKeys_df_dummy_list = []
    db = cluster["DRI1"]
    date_from_SI = ""
    filter = False
    if "date_from_SI" in session:
        date_from_SI = session.pop("date_from_SI",None)
    form_type = "DRI 1"
    if str(session["formation"]) == "DRI HQ" or str(session["formation"])=="Admin":
        filter = False
        total,doc,allKeys_df_dummy_list,allkeydataheadings =vi_col_dri1(list(db.list_collection_names()))
    else:
        if date_from_SI!="":
            filter = True
            doc,allKeys_df_dummy_list,total,allkeydataheadings = filter_by_date_dri1(db,session["formation"],form_type,date_from_SI)
        else:
            filter = False
            total,doc,allKeys_df_dummy_list,allkeydataheadings=vi_col_dri1([str(session["formation"])])
    page = request.args.get('page', 1,type=int)
    ROWS_PER_PAGE = 4
    limit = 1
    print(total)
    offset = (int(page)) * limit-limit
    pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
    print(allKeys_df_dummy_list)
    headings=[]
    for val in list(doc.keys()):
        headings.append(str(val).capitalize())
    if doc:
        headings.sort()
        temp = headings[0]
        headings[0] = headings[1]
        headings[1] = temp
    dri_db = cluster["DRI"]
    if "formations" in list(dri_db.list_collection_names()):
        collection = dri_db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        zonal_units=formation_data.values.tolist()
    else:
        zonal_units=[]
    return render_template("view_DRI1.html",content=navbar_conditions(),zonal_units=zonal_units, headings = headings, pagination=pagination, allkeysdata = allKeys_df_dummy_list,allkeydataheadings=allkeydataheadings,filter=filter)

@app.route('/clear_filter',methods=["POST","GET"])
def clear_filter():
    if request.method == "POST":
        button_req = request.form.get("button_req")
        if button_req == "Inputs Shared With Your Formation":
            return redirect(url_for("view_intelligence"))
        if button_req == "DRI 1":
            return redirect(url_for("view_DRI1"))
        if button_req == "DRI 2":
            return redirect(url_for("view_DRI2"))

def vi_col_dri2(col_name_list):
    db = cluster["DRI2"]
    total_data_final = pd.DataFrame()
    total = 0
    doc = {}
    place_list = []
    for val in col_name_list:
        collection1 = db[str(val)]
        print("here")
        placedropdown = collection1.find({"Form": "DRI 2"}, {'place': True})
        for i in placedropdown:
            place_list.append(i['place'])
            place_list.sort()
        place_data = set(place_list)
        dri1_allKeys = collection1.find({"Form": "DRI 2"}, {'_id':False, 'serialno':True,'date':True,  'caseno': True ,'commissionarate': True, 'division': True })
        dri1_allKeys_df = pd.DataFrame(list(collection1.find({"Form": "DRI 2"}, {'_id':False, 'serialno':True, 'date':True,  'caseno': True ,'commissionarate': True, 'division': True })))
        total = total + len(dri1_allKeys_df)
        allKeys_df_dummy = pd.DataFrame(list(collection1.find({"Form": "DRI 2"}, { 'serialno':True, 'date':True,  'caseno': True ,'commissionarate': True, 'division': True })))
        allKeys_df_dummy["collection_id"] = id
        total_data_final = pd.concat([total_data_final, allKeys_df_dummy], ignore_index=True)
        for doc in dri1_allKeys:
             pass
        allKeys_df_dummy_list = total_data_final.to_dict(orient="records")
        return place_data,total,doc,allKeys_df_dummy_list,list(allKeys_df_dummy)

@app.route("/view-DRI2")
def view_DRI2():
    allKeys_df_dummy_list = []
    db = cluster["DRI2"]
    print(session["formation"])
    print("helloooooooooooooooooooooooooooooooooooo")
    if str(session["formation"]) == "DRI HQ" or str(session["formation"]) == "Admin":
        place_data,total,doc,allKeys_df_dummy_list,allkeydataheadings =vi_col_dri2(list(db.list_collection_names()))
        print(db.list_collection_names())
    else:
        place_data,total,doc,allKeys_df_dummy_list,allkeydataheadings=vi_col_dri2([str(session["formation"])])
    page = request.args.get('page', 1)
    ROWS_PER_PAGE = 4
    limit = 10
    offset = (int(page)) * limit-limit
    pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
    print(allKeys_df_dummy_list)
    headings=[]
    for val in list(doc.keys()):
        headings.append(str(val).capitalize())
    if doc:
        headings.sort()
        temp = headings[0]
        headings[0] = headings[1]
        headings[1] = temp
    return render_template("view_DRI2.html",content=navbar_conditions(),AU=False, SI=False, filter_place= place_data, headings = headings, pagination=pagination, allkeysdata = allKeys_df_dummy_list,allkeydataheadings=allkeydataheadings)
    

def place_data_fun_with(collection1):
    place_list=[]
    placedropdown = collection1.find({"Form": "Share inputs with your formation"})
    for i in placedropdown:
        print(i)
        place_list.append(i['place'])
        place_list.sort()
    place_data = set(place_list)
    print(place_data)
    return place_data

def place_data_fun_by(collection1):
    place_list=[]
    placedropdown = collection1.find({"Form": "Share inputs by your formation"})
    for i in placedropdown:
        print(i)
        place_list.append(i['place'])
        place_list.sort()
    place_data = set(place_list)
    print(place_data)
    return place_data


def vi_col(col_name_list):
    db = cluster["Share_Intelligence"]
    total_data_final = pd.DataFrame()
    total = 0
    doc = {}
    for val in col_name_list:
        collection1 = db[str(val)]
        place_data = place_data_fun_with(collection1)
        print(place_data,"in view intelligence")
        # source_data = collection1.find({"Form": "Share inputs by your formation"}, {'id': False})
        # for i in source_data:
        #     print(i)
        allKeys = collection1.find({"Form": "Share inputs with your formation"}, {'_id': False,'deputy_director_email_id': False, 'place':False,'Estimated Value(Value in Crores-Rs)':False,
        'Suspect name':False,'unit':False,'Form': False,'login_link':False,'recipients':False, 'gist_of_input': False, 'additional_director_email_id': False,"other_commodity":False,  'zonal_head_email_id':False, 'file':False, 'priority':False, 'Upload Supporitng Documents':False })
        allKeys_df = pd.DataFrame(list(collection1.find({"Form": "Share inputs with your formation"}, {'_id': False,'deputy_director_email_id': False, 
        'Estimated Value(Value in Crores-Rs)':False, 'Form': False,'place':False,'Suspect name':False, 'additional_director_email_id': False, 'deputy_director_email_id': False, 'zonal_head_email_id':False, 'file':False, 'priority':False , 'Upload Supporitng Documents':False})))
        total = total + len(allKeys_df)
        allKeys_df_dummy = pd.DataFrame(list(collection1.find({"Form": "Share inputs with your formation"}, {'additional_director_email_id': False,'deputy_director_email_id': False,'Estimated Value(Value in Crores-Rs)':False,'Suspect name':False,'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
        total_data_final = pd.concat([total_data_final, allKeys_df_dummy], ignore_index=True)
        for doc in allKeys:
            pass
    allKeys_df_dummy_list = total_data_final.to_dict(orient="records")
    return place_data,len(total_data_final),doc,allKeys_df_dummy_list,list(allKeys_df_dummy)
    

@app.route("/view-intelligence")
def view_intelligence():
    allKeys_df_dummy_list = []
    db = cluster["Share_Intelligence"]
    place_value = ""
    doc={}
    date_from_SI = ""
    filter = False
    if "place_name_from_SI" in session:
        place_value = session.pop("place_name_from_SI",None)
    if "date_from_SI" in session:
        date_from_SI = session.pop("date_from_SI",None)
    form_type = "Share inputs with your formation"
    if str(session["formation"]) == "DRI HQ" or str(session["formation"])=="Admin":
        filter = False
        place_data,total,doc,allKeys_df_dummy_list,allkeydataheadings =vi_col(list(db.list_collection_names()))
    else:
        if place_value!="" and date_from_SI=="":
            filter = True
            place_data,doc,allKeys_df_dummy_list,total,allkeydataheadings = filter_by_place(db,place_value,session["formation"],form_type)
        elif date_from_SI!="":
            if place_value=="":
                filter = True
                place_data,doc,allKeys_df_dummy_list,total,allkeydataheadings = filter_by_date(db,session["formation"],form_type,date_from_SI)
            elif place_value!="":
                filter = True
                place_data,doc,allKeys_df_dummy_list,total,allkeydataheadings = filter_by_date_and_place(db,place_value,session["formation"],form_type,date_from_SI)
        else:
            filter = False
            place_data,total,doc,allKeys_df_dummy_list,allkeydataheadings=vi_col([str(session["formation"])])
    page = request.args.get('page', 1,type=int)
    ROWS_PER_PAGE = 4
    limit = 1
    print(total)
    offset = (int(page)) * limit-limit
    pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
    print(allKeys_df_dummy_list)
    headings=[]
    for val in list(doc.keys()):
        headings.append(str(val).capitalize())
    if doc:
        headings.sort()
        temp = headings[0]
        headings[0] = headings[1]
        headings[1] = temp
        headings.remove("Regional_unit")
        headings.remove("Sub_regional_unit")
    dri_db = cluster["DRI"]
    if "formations" in list(dri_db.list_collection_names()):
        collection = dri_db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        zonal_units=formation_data.values.tolist()
    else:
        zonal_units=[]
    print(place_data,"In share intel view rout")
    return render_template("view_intelligence.html",content=navbar_conditions(),zonal_units=zonal_units,filter_place= place_data, headings = headings, pagination=pagination, allkeysdata = allKeys_df_dummy_list,allkeydataheadings=allkeydataheadings,filter=filter)


@app.route('/download',  methods=['GET', 'POST'])
def download():
    if request.method == "POST":
        db_name = request.form.get("db_name")
        db = cluster[str(db_name)]
        collection_id = request.form.get("collection_id")
        id_value = request.form.get("id_value")
        input_data = db[str(collection_id)]
        id_value = ObjectId(id_value)
        Form = request.form.get("Form")
        try:
            if input_data.find({"_id":id_value}):
                pandas_data_frame = pd.DataFrame(input_data.find({"_id":id_value},{'Upload Supporitng Documents':False,'additional_director_email_id': False,'deputy_director_email_id': False,'Estimated Value(Value in Crores-Rs)':False,'Suspect name':False,'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False }),index=[0])
                # print(pandas_data_frame)
                pandas_data_frame["_id"][0] = str(pandas_data_frame["_id"][0])
                import docx
                doc = docx.Document()
                for col in pandas_data_frame.columns:
                    if col!="_id":
                        if not pandas_data_frame[col][0] == "":
                            temp = str(col).replace("_"," ")
                            doc.add_heading(str(temp).capitalize(), 1)
                            doc.add_paragraph(str(pandas_data_frame[col][0]))
                        else:
                            temp = str(col).replace("_"," ")
                            doc.add_heading(str(temp).capitalize(), 1)
                            doc.add_paragraph("None")
                doc.save('download.docx')
                return send_file('download.docx', download_name="download.docx", as_attachment=True )
            
        except Exception as e:
                print(e)
                if Form == "Share inputs by your formation":
                    return redirect(url_for("view-intelligence2"))
                elif Form == "Share inputs with your formation":
                    return redirect(url_for("view_intelligence"))
                elif Form == "DRI 1":
                    return redirect(url_for("view_DRI1"))
                else:
                    return redirect(url_for("view_DRI2"))

@app.route('/download_document', methods=['GET', 'POST'])
def download_document():
    try:
        print("I am in try of dd")
        if not os.path.exists(app.root_path + "/download_documents"):
            os.mkdir(app.root_path + "/download_documents")
        for previous_input in os.listdir(app.root_path + "/download_documents"):
            os.remove(app.root_path + "/download_documents/" + previous_input)
       
        file_id_list = []
        
        if request.method == "POST":
            db_name = request.form.get("db_name")
            db = cluster[str(db_name)]
            collection_id = request.form.get("collection_id")
            print(collection_id,"collection_id")
            id_value = request.form.get("id_value")
            id_value = ObjectId(id_value)
            input_data = db[collection_id]
            out = input_data.find({"_id": id_value})
            for i in out:
                file_id_list = i["Upload Supporitng Documents"]
            for file_id in file_id_list:
                fs = gridfs.GridFS(db)
                file = fs.get(list(file_id.values())[0])
                print(file)
                print("-------------------------------------------")
                print(file_id,"-----------------------")
                with open("download_documents/" + list(file_id.keys())[0], "wb") as fd:
                    fd.write(file.read())

            name = 'download_documents'
            zip_name = name + '.zip'
            with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for folder_name, subfolders, filenames in os.walk(name):
                    zip_ref.write(folder_name)
                    for filename in filenames:
                        zip_ref.write(folder_name + '/' + filename)

            zip_ref.close()

            path = send_file(app.root_path + "/download_documents.zip", as_attachment=True)
            return path

    except Exception as e:
        print(e)
        name = 'download_documents'
        zip_name = name + '.zip'
        with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for folder_name, subfolders, filenames in os.walk(name):
                zip_ref.write(folder_name)
                for filename in filenames:
                    zip_ref.write(folder_name + '/' + filename)

        path = send_file(app.root_path + "/download_documents.zip", as_attachment=True)
        return path

def filter_by_place(db,place,formation,form_type):
    collection1 = db[formation]
    total = 0
    doc={}
    if "with" in form_type:
        place_data = place_data_fun_with(collection1)
    if "by" in form_type:
        place_data = place_data_fun_by(collection1)
    print(place_data,"In filter by place fun")
    dri1_allKeys = collection1.find({"$and": [{"Form": f"{form_type}"},{"place": place }]}, {'_id': False,'deputy_director_email_id': False, 'place':False,'Estimated Value(Value in Crores-Rs)':False,
        'Suspect name':False,'unit':False,'Form': False,'login_link':False,'recipients':False, 'gist_of_input': False, 'additional_director_email_id': False,"other_commodity":False,  'zonal_head_email_id':False, 'file':False, 'priority':False, 'Upload Supporitng Documents':False })
    dri1_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"place": place }]}, {'_id': False, 'additional_director_email_id': False,'place':False , 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    total = total + len(dri1_allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"place": place }]}, {'additional_director_email_id': False, 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    for doc in dri1_allKeys:
        pass
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    return place_data,doc,allKeys_df_dummy_list,total,list(allKeys_df_dummy)
def filter_by_date(db,formation,form_type,datevalue):
    collection1 = db[formation]
    total=0
    doc={}
    place_data = place_data_fun_with(collection1)
    print(datevalue)
    # date_time_obj = datetime.strptime(datevalue, '%y-%m-%d')
    print(place_data," in filter by date")
    dri1_allKeys = collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]}, {'_id': False, 'formation':False,'deputy_director_email_id':False,'additional_director_email_id': False,'place':False , 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False, 'Upload Supporitng Documents':False })
    dri1_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]}, {'_id': False,'deputy_director_email_id':False, 'additional_director_email_id': False,'place':False , 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    total = total + len(dri1_allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]}, {'deputy_director_email_id':False,'additional_director_email_id': False, 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    for doc in dri1_allKeys:
        pass
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    return place_data,doc,allKeys_df_dummy_list,total,list(allKeys_df_dummy)
def filter_by_date_and_place(db,place,formation,form_type,datevalue):
    collection1 = db[formation]
    total = 0
    doc={}
    place_data = place_data_fun_with(collection1)
    print(place_data," in filter by date")
    dri1_allKeys = collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue},{"place": place }]}, {'_id': False, 'formation':False,'deputy_director_email_id':False,	'additional_director_email_id': False,'place':False , 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False, 'Upload Supporitng Documents':False })
    dri1_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue},{"place": place }]}, {'_id': False, 'deputy_director_email_id':False,'additional_director_email_id': False,'place':False , 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    total = total + len(dri1_allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue},{"place": place }]}, {'deputy_director_email_id':False,'additional_director_email_id': False, 'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    for doc in dri1_allKeys:
        pass
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    return place_data,doc,allKeys_df_dummy_list,total,list(allKeys_df_dummy)

@app.route('/shareinputs_filters',  methods=['GET', 'POST'])
def shareinputs_filters():
    doc = {}
    if request.method == "POST":
        db = cluster["Share_Intelligence"]
        list_place = []
        total = 0
        place_name = request.form.get("place_name")
        print(place_name)
        datevalue = request.form.get('date')
        form_type = request.form.get('form_type')
        doc={}
        for col in list(db.list_collection_names()):
            if str(col)==str(session["formation"]):
                if str(place_name)!="" and str(datevalue)=="":
                    session["place_name_from_SI"] = place_name
                    # place_data,doc,allKeys_df_dummy_list,total = filter_by_place(db,place_name,session["formation"],form_type)
                if str(datevalue)!="":
                    if str(place_name)!="":
                        session["place_name_from_SI"] = place_name
                        session["date_from_SI"] = datevalue
                        # place_data,doc,allKeys_df_dummy_list,total = filter_by_date_and_place(db,place_name,session["formation"],form_type,datevalue)
                    if str(place_name)=="":
                        session["date_from_SI"] = datevalue
                        # place_data,doc,allKeys_df_dummy_list,total = filter_by_date(db,session["formation"],form_type,datevalue)
        if "with" in form_type:
            return redirect(url_for("view_intelligence"))
        if "by" in form_type:
            return redirect(url_for("view_intelligence2"))
    
    
   


# *********** DRI1 filters ***********





def filter_by_date_dri1(db,formation,form_type,datevalue):
    collection1 = db[formation]
    total=0
    print(datevalue)
    # date_time_obj = datetime.strptime(datevalue, '%y-%m-%d')
    dri1_allKeys = collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]},  {'_id': False,  'Form': False, 'place':False,'other_commodity':False, 'involve': False, 'links': False, 'action':False, 'remarks':False, 'name':False, 'designation':False,'DRI Zonal Unit':False, 'Upload Supporitng Documents': False })
    dri1_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]},{'_id': False,'place':False,'other_commodity':False,  'Form': False, 'involve': False, 'links': False, "place": False,'action':False, 'remarks':False, 'name':False, 'designation':False,'DRI Zonal Unit':False, 'Upload Supporitng Documents': False })))
    total = total + len(dri1_allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]}, {'Form': False, 'involve': False, 'links': False, 'action':False, 'remarks':False, 'name':False, 'designation':False,'DRI Zonal Unit':False, 'Upload Supporitng Documents': False })))
    for doc in dri1_allKeys:
        pass
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    return doc,allKeys_df_dummy_list,total,list(allKeys_df_dummy)         
        
@app.route('/dri1_filters',  methods=['GET', 'POST'])
def dri1_filters():
    if request.method == "POST":
        doc = {}
        allKeys_df_dummy_list = []
        db = cluster["DRI1"]
        list_place = []
        total = 0
        datevalue = request.form.get('date')
        form_type = request.form.get('form_type')
        doc={}
        for col in list(db.list_collection_names()):
            if str(col)==str(session["formation"]):
                if str(datevalue)!="":
                    session["date_from_SI"] = datevalue
                        # place_data,doc,allKeys_df_dummy_list,total = filter_by_date(db,session["formation"],form_type,datevalue)
        return redirect(url_for("view_DRI1"))
    






    # DRI2 filters
def filter_by_date_dri2(db,formation,form_type,datevalue):
    collection2 = db[formation]
    total=0
    print(datevalue)
    # date_time_obj = datetime.strptime(datevalue, '%y-%m-%d')
    dri1_allKeys = collection2.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]},   {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })
    dri1_allKeys_df = pd.DataFrame(list(collection2.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]},  {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })))
    total = total + len(dri1_allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection2.find({"$and": [{"Form": f"{form_type}"},{"date": datevalue}]},  {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })))
    for doc in dri1_allKeys:
        pass
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    return doc,allKeys_df_dummy_list,total,list(allKeys_df_dummy)         
        
@app.route('/dri2_filters',  methods=['GET', 'POST'])
def dri2_filters():
    if request.method == "POST":
        doc = {}
        allKeys_df_dummy_list = []
        db = cluster["DRI2"]
        list_place = []
        total = 0
        datevalue = request.form.get('date')
        form_type = request.form.get('form_type')
        doc={}
        for col in list(db.list_collection_names()):
            if str(col)==str(session["formation"]):
                if str(datevalue)!="":
                    session["date_from_SI"] = datevalue
                        # place_data,doc,allKeys_df_dummy_list,total = filter_by_date(db,session["formation"],form_type,datevalue)
        return redirect(url_for("view_DRI2"))
    
        

# @app.route('/dri2_filters',  methods=['GET', 'POST'])
# def dri2_filters(): 
#     collection1 = db[session["uid"]]
#     if request.method == "POST":
#         try:
#             datevalue = request.form.get('date')
#             place_name = request.form.get('place_name')
#             placedropdown = collection1.find({"Form": "DRI 2"}, {'place': True})
#             place_list_data = []
#             for i in placedropdown:
#                 place_list_data.append(i['place'])
#                 place_list_data.sort()
#             place_data = set(place_list_data)

#             session["place_data"]=list(place_data)
#             session["datevalue"]=datevalue
#             session["place_name"]=place_name
#             dri2_allKeys = collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#             {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })
#             dri2_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#             {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })))

#             # Set the pagination configuration
#             page = request.args.get('page',1)
#             # page = request.args.get('page', 1, type=int)
#             limit = 10
#             ROWS_PER_PAGE = 4
#             offset = (int(page)) * limit-limit
#             total = len(dri2_allKeys_df)
#             allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#             { 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True }).skip(offset).limit(limit)))
#             pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
#             for doc in dri2_allKeys:
#                 pass
            
#             return render_template("view_DRI2.html", dri2_headings = doc.keys(), pagination=pagination, allkeysdata = allKeys_df_dummy.values.tolist(),filter_place = list(place_data))
#         except:
#             msg="You may have empty data or Please give all the required inputs"
#             return render_template("view_DRI2.html", msg=msg, pagination=pagination, allkeysdata = allKeys_df_dummy.values.tolist(),filter_place = session["place_data"])
    
#     else:
#         dri2_allKeys = collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#         {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })
#         dri2_allKeys_df = pd.DataFrame(list(collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#         {'_id':False, 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True })))

#         # Set the pagination configuration
#         page = request.args.get('page',1)
#         limit = 10
#         ROWS_PER_PAGE = 4
#         offset = (int(page)) * limit-limit
#         total = len(dri2_allKeys_df)
#         allKeys_df_dummy = pd.DataFrame(list(collection1.find({"$and": [{"Form": "DRI 2"},{"date": session["datevalue"] },{"place": session["place_name"] }]},
#         { 'serialno': True, 'date': True, 'place': True, 'caseno': True ,'commissionarate': True, 'division': True }).skip(offset).limit(limit)))
#         pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
#         for doc in dri2_allKeys:
#             pass

#         return render_template("view_DRI2.html", dri2_headings = doc.keys(), pagination=pagination, allkeysdata = allKeys_df_dummy.values.tolist(),filter_place = session["place_data"])


@app.route("/dri1form", methods=['POST'])
def dri1form():
    if request.method == "POST":
        try:
            commodity = request.form.get('commodities')
            selected_commodity = request.form.get('commodities')
            other_commodity = ""
            if selected_commodity == 'others':
                 other_commodity = request.form.get('Commodities')
            value = request.form.get('value')
            suspect = request.form.get('suspect')
            involve = request.form.get('involve')
            links = request.form.get('links')
            place = request.form.get('place')
            date = request.form.get('date')
            action = request.form.get('action')
            remarks = request.form.get('remarks')
            officer = request.form.get('officer')
            designation = request.form.get('designation')
            DriZonalUnit = request.form.get('DriZonalUnit')
            file = request.files.getlist('files')
            files_list = list()
            db = cluster["DRI1"]
            branch_name = str(session["formation"])
            collection = db[branch_name]
            print(collection)
            fs = gridfs.GridFS(db)
            for fi in file:
                fileID = fs.put(fi.stream.read())
                print("file_id", fileID)
                files_list.append({fi.filename : fileID})
            dri_form = {"Form":"DRI 1","commodity":commodity,"other_commodity":other_commodity,"Value":value,"Suspect":suspect,"involve":involve,"involve":involve,"links":links,
                "place":place, "date":date, "action":action,"remarks":remarks,"officer":officer,"designation":designation,"DRI Zonal Unit":DriZonalUnit, "Upload Supporitng Documents": files_list}
            # Insert Data
            collection.insert_one(dri_form)
            pandas_data_frame = pd.DataFrame(dri_form)
            pandas_data_frame["_id"][0] = str(pandas_data_frame["_id"][0])
            import docx
            doc = docx.Document()
            for col in pandas_data_frame.columns:
                if not pandas_data_frame[col][0] == "":
                    doc.add_heading(col, 3)
                    doc.add_paragraph(str(pandas_data_frame[col][0]))
                else:
                    doc.add_heading(col, 3)
                    doc.add_paragraph("")
            doc.save(r'mail/New_DRI_1.docx')
            msg = Message(subject='New DRI-1 Document Generated', sender='isra.government@gmail.com', recipients=["manohari.manu135@gmail.com"])
            msg.body = str("Please find the attachment for the details of newly generated DRI-1 Document")
            with app.open_resource(r'mail/New_DRI_1.docx') as fp:
                msg.attach(r'New_DRI_1.docx', "text/docx", fp.read())
            mail.send(msg)
            print(session["email_id"])
            session["dri1_filled_msg"]="Your form has been submitted"
            return redirect(url_for("create_dri1"))
        except Exception as e:
            print(e)
            session["error_msg"] = f"{e}"
            return redirect(url_for("create_dri1"))
        

@app.route("/forward_button", methods=['GET', 'POST'])
def forward_button():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        id_value = ObjectId(id_value)
        collection_id = request.form.get("collection_id")
        input_data = db[collection_id]
        Form = request.form.get("Form")
        try:
            if input_data.find({"_id":id_value}):
                pandas_data_frame = pd.DataFrame(input_data.find({"_id":id_value}),index=[0])
                pandas_data_frame["_id"][0] = str(pandas_data_frame["_id"][0])
                import docx
                doc = docx.Document()
                for col in pandas_data_frame.columns:
                    if not pandas_data_frame[col][0] == "":
                        doc.add_heading(col, 3)
                        doc.add_paragraph(str(pandas_data_frame[col][0]))
                    else:
                        doc.add_heading(col, 3)
                        doc.add_paragraph("")
                doc.save('mail/forward_DRI1.docx')
            msg = Message(subject='New DRI-1 Document Generated', sender='isra.government@gmail.com', recipients=[session["email_id"]])
            msg.body = f"""Hi,

                        Please find the attachment for the details of {id_value} DRI-1 Document below"""
            with app.open_resource(r'mail/forward_DRI1.docx') as fp:
                msg.attach(r'forward_DRI1.docx', "text/docx", fp.read())
            mail.send(msg)
            print(session["email_id"])
            return render_template("home.html", msg="Your form has been submitted")
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        





@app.route("/dri2form", methods=['POST'])
def dri2form():
    
    if request.method == "POST":
        db = cluster["DRI2"]
        branch_name = str(session["formation"])
        collection = db[branch_name]
        print(collection)
        serialno = request.form.get('serialno')
        date = request.form.get('date')
        place = request.form.get('place')
        caseno = request.form.get('caseno')
        information = request.form.get('information')
        fileno = request.form.get('fileno')
        others = request.form.get('others')
        dateofdetection = request.form.get('dateofdetection')
        commissionarate = request.form.get('commissionarate')
        division = request.form.get('division')
        trafficking = request.form.get('trafficking')
        seizure = request.form.get('seizure')
        amountofduty = request.form.get('amountofduty')
        importduty = request.form.get('importduty')
        export = request.form.get('export')
        antidumping = request.form.get('antidumping')
        fobvalue = request.form.get('fobvalue')
        desription = request.form.get('description')
        modus = request.form.get('modus')
        goods_involved = request.form.get('goods_involved')
        cth = request.form.get('cth')
        country = request.form.get('country')
        date_of_seizure = request.form.get('date_of_seizure')
        time_of_seizure = request.form.get('time_of_seizure')
        place_of_seizure = request.form.get('place_of_seizure')
        other_places = request.form.get('other_places')
        goods_seized = request.form.get('goods_seized')
        cth2 = request.form.get('cth2')
        quantity = request.form.get('quantity')
        unit = request.form.get('unit')
        marking = request.form.get('marking')
        brandname = request.form.get('brandname')
        model = request.form.get('model')
        make = request.form.get('make')
        type_of_packages = request.form.get('type_of_packages')
        value_of_conveyance = request.form.get('value_of_conveyance')
        toatal_value = request.form.get('toatal_value')
        total_value_of_seizure = request.form.get('total_value_of_seizure')
        name_of_firm = request.form.get('name_of_firm')
        nature_of_firm = request.form.get('nature_of_firm')
        name_of_proprietor = request.form.get('name_of_proprietor')
        address = request.form.get('address')
        telephone = request.form.get('telephone')
        address_firm_abroad = request.form.get('address_firm_abroad')
        telephone_number = request.form.get('telephone_number')
        business = request.form.get('business')
        trader = request.form.get('trader')
        others_ = request.form.get('others_')
        iec = request.form.get('iec')
        pan_no = request.form.get('pan_no')
        pan_date = request.form.get('pan_date')
        pan_place = request.form.get('pan_place')
        address_of_factory = request.form.get('address_of_factory')
        address_of_godown = request.form.get('address_of_godown')
        account_no = request.form.get('account_no')
        offences = request.form.get('offences')
        person_name = request.form.get('person_name')
        alias = request.form.get('alias')
        father_name = request.form.get('father_name')
        nationality = request.form.get('nationality')
        date_of_birth = request.form.get('date_of_birth')
        status = request.form.get('status')
        residential_address = request.form.get('residential_address')
        residential_phone_numbers = request.form.get('residential_phone_numbers')
        mobile_numbers = request.form.get('mobile_numbers')
        profession = request.form.get('profession')
        abroad_numbers = request.form.get('abroad_numbers')
        aircraft = request.form.get('aircraft')
        details_aircraft = request.form.get('details_aircraft')
        details_passport = request.form.get('details_passport')
        pp = request.form.get('pp')
        pp_dl = request.form.get('pp_dl')
        pp_dl_elec = request.form.get('pp_dl_elec')
        pan_no = request.form.get('pan_no')
        pan_date_of_issue = request.form.get('pan_date_of_issue')
        pan_place_of_issue = request.form.get('pan_place_of_issue')
        bank_name = request.form.get('bank_name')
        ac_no = request.form.get('ac_no')
        interogated = request.form.get('interogated')
        arrested = request.form.get('arrested')
        offences = request.form.get('offences')
        convictions = request.form.get('convictions')
        cofeposa = request.form.get('cofeposa')
        name_of_cha = request.form.get('name_of_cha')
        name_partners = request.form.get('name_partners')
        address_of_cha = request.form.get('address_of_cha')
        tele = request.form.get('tele')
        cha = request.form.get('cha')
        custom = request.form.get('custom')
        key_person = request.form.get('key_person')
        prev_offences = request.form.get('prev_offences')
        no_of_be = request.form.get('no_of_be')
        date_of_be = request.form.get('date_of_be')
        consignor_name = request.form.get('consignor_name')
        consignor_address = request.form.get('consignor_address')
        consignee_name = request.form.get('consignee_name')
        consignee_address = request.form.get('consignee_address')
        name_of_overseas_buyer = request.form.get('name_of_overseas_buyer')
        address_of_overseas = request.form.get('address_of_overseas')
        depart = request.form.get('depart')

        dri_form = {"Form":"DRI 2","serialno":serialno,"date":date,"place":place,"caseno":caseno,"information":information,"fileno":fileno,
        "others":others,"dateofdetection":dateofdetection,
        "commissionarate":commissionarate,"division":division,"trafficking":trafficking,"seizure":seizure,
        "amountofduty":amountofduty,"importduty":importduty,
        "export":export,"antidumping":antidumping,"fobvalue":fobvalue,"description":desription,"modus":modus,"goods_involved":goods_involved,"cth":cth,
        "country":country,"date_of_seizure":date_of_seizure,"time_of_seizure":time_of_seizure,"place_of_seizure":place_of_seizure,
        "other_places":other_places,"goods_seized":goods_seized,"cth2":cth2,"quantity":quantity,"unit":unit,"marking":marking,"brandname":brandname,
        "model":model,"make":make,"type_of_packages":type_of_packages,"value_of_conveyance":value_of_conveyance,"toatal_value":toatal_value,"total_value_of_seizure":total_value_of_seizure,"name_of_firm":name_of_firm,"nature_of_firm":nature_of_firm,"name_of_proprietor":name_of_proprietor,"address":address,"telephone":telephone,"address_firm_abroad":address_firm_abroad,"telephone_number":telephone_number,
        "business":business,"trader":trader,"others_":others_,"iec":iec,"pan_no":pan_no,"pan_date":pan_date,"pan_place":pan_place,"address_of_factory":address_of_factory,
        "address_of_godown":address_of_godown,"account_no":account_no,"offences":offences,"person_name":person_name,"alias":alias,"father_name":father_name,
        "nationality":nationality,"date_of_birth":date_of_birth,"status":status,"residential_address":residential_address,"residential_phone_numbers":residential_phone_numbers,
        "mobile_numbers":mobile_numbers,"profession":profession,"abroad_numbers":abroad_numbers,"aircraft":aircraft,"details_aircraft":details_aircraft,
        "details_passport":details_passport,"pp":pp,"pp_dl":pp_dl,"pp_dl_elec":pp_dl_elec,"pan_no":pan_no,"pan_date_of_issue":pan_date_of_issue,
        "pan_place_of_issue":pan_place_of_issue,"bank_name":bank_name,"ac_no":ac_no,"interogated":interogated,"arrested":arrested,
        "offences":offences,"convictions":convictions,"cofeposa":cofeposa,"name_of_cha":name_of_cha,"name_partners":name_partners,"address_of_cha":address_of_cha,"tele":tele,
        "cha":cha,"custom":custom,"key_person":key_person,"prev_offences":prev_offences,"no_of_be":no_of_be,"date_of_be":date_of_be,"consignor_name":consignor_name,
        "consignor_address":consignor_address,"consignee_name":consignee_name,"consignee_address":consignee_address,"name_of_overseas_buyer":name_of_overseas_buyer,
        "address_of_overseas":address_of_overseas,"depart":depart}

        # Insert Data
        collection.insert_one(dri_form)

        msg = "Your form has been submitted"
        return render_template("create_new_DRI2.html",content=navbar_conditions(),form_type={"formtype":branch_name}, msg=msg)




@app.route("/shareintelligence-form-post", methods=['POST'])
def shareintelligenceform():
    if request.method == "POST":
        db = cluster["Share_Intelligence"]
        branch_name = str(session["formation"])
        collection = db[branch_name]
        print(collection)
        date = request.form.get('date')
        place = request.form.get('place')
        gist_of_input = request.form.get('gist_of_input')
        # unit = request.form.get('unit')
        regional_Unit = request.form.get('regional_Unit')
        sub_regional_unit = request.form.get('sub_regional_unit')
        commodity = request.form.get('Commodity')
        selected_commodity = request.form.get('Commodity')
        other_commodity = ""


        if selected_commodity == 'others':
            other_commodity = request.form.get('commodity')
        
        
        intel_reciever_email_id = request.form.get('intel_reciever_email_id')
        deputy_director_email_id = request.form.get('deputy_director_email_id')

        zonal_head_email_id = request.form.get('zonal_head_email_id')
        additional_director_email_id = request.form.get('additional_director_email_id')
        priority = request.form.get('priority')
        file = request.files.getlist('files')
        login_link = request.url_root + 'login'
        recipients = [request.form.get('intel_reciever_email_id'), request.form.get('zonal_head_email_id'),request.form.get('additional_director_email_id'),request.form.get('deputy_director_email_id')]

        print(login_link)

        # form_link = request.url_root
        files_list = list()
        fs = gridfs.GridFS(db)
        for fi in file:
            fileID = fs.put(fi.stream.read())
            print("file_id", fileID)
            files_list.append({fi.filename : fileID})
        share_intel_form = {"Form": "Share inputs by your formation", "formation": branch_name, "date": date, "place": place,
                            "gist_of_input": gist_of_input, "regional_Unit": regional_Unit,
                            "sub_regional_unit": sub_regional_unit, "commodity": commodity, "other_commodity":other_commodity,
                            "intel_reciever_email_id": intel_reciever_email_id,"deputy_director_email_id":deputy_director_email_id,
                            "zonal_head_email_id": zonal_head_email_id,
                            "additional_director_email_id": additional_director_email_id,
                            "priority": priority, "Upload Supporitng Documents": files_list}
        collection.insert_one(share_intel_form)
        collection2 = db[str(place)]
        share_intel_form2 = share_intel_form.copy()
        share_intel_form2["Form"] = "Share inputs with your formation"
        share_intel_form2["formation"] = share_intel_form["place"]
        share_intel_form2["place"] = share_intel_form["formation"] 
        collection2.insert_one(share_intel_form2)
        pandas_data_frame = pd.DataFrame(share_intel_form)
        pandas_data_frame["_id"][0] = str(pandas_data_frame["_id"][0])
        import docx
        doc = docx.Document()
        for col in pandas_data_frame.columns:
            if not pandas_data_frame[col][0] == "":
                doc.add_heading(col, 3)
                doc.add_paragraph(str(pandas_data_frame[col][0]))
            else:
                doc.add_heading(col, 3)
                doc.add_paragraph("")
        doc.save(r'mail/new_Intelshared.docx')
        for recipient in recipients: 
            msg = Message(subject=f"Intel shared from {place}", sender='isra.government@gmail.com',recipients=[recipient])
            msg.body = str(f"Please note that an intelligence is shared by {branch_name} with {place}. Click Here to Login {login_link}")
            # msg.body += f" Click <a href='{login_link}'>here</a> to access the login page."

            with app.open_resource(r'mail/new_Intelshared.docx') as fp:
                msg.attach(r'new_Intelshared.docx', "text/docx", fp.read())
            # form_link = request.url_root
            mail.send(msg)

        msg = "Your form has been submitted"
        return render_template("shareintelligence2.html",content=navbar_conditions(), buttons=button_conditions(), msg=msg, form_type={"formtype":branch_name})


@app.route("/intelligence_view_type",methods=["POST"])
def intelligence_view_type():
    if request.method == "POST":
        button_req = request.form.get("button_req")
        if button_req == "Inputs Shared With Your Formation":
            return redirect(url_for("view_intelligence"))
        elif button_req == "Inputs Shared By Your Formation":
            return redirect(url_for("view_intelligence2"))

@app.route("/view-intelligence2")
def view_intelligence2():
    doc = {}
    allKeys_df_dummy_list = []
    place_list = []
    total = 0
    db = cluster["Share_Intelligence"]
    collection1 = db[str(session["formation"])]
    placedropdown = collection1.find({"Form": "Share inputs by your formation"}, {'place': True})
    for i in placedropdown:
        place_list.append(i['place'])
        place_list.sort()
    place_data = set(place_list)
    allKeys = collection1.find({"Form": "Share inputs by your formation"}, {'_id': False,'deputy_director_email_id': False, 'place':False,'Estimated Value(Value in Crores-Rs)':False,"other_commodity":False,
    'Suspect name':False,'unit':False,'Form': False,'login_link':False,'recipients':False, 'gist_of_input': False, 'additional_director_email_id': False,  'zonal_head_email_id':False, 'file':False, 'priority':False, 'Upload Supporitng Documents':False })
    allKeys_df = pd.DataFrame(list(collection1.find({"Form": "Share inputs by your formation"}, {'_id': False,'deputy_director_email_id': False, 
    'Estimated Value(Value in Crores-Rs)':False, 'Form': False,'place':False,'Suspect name':False, 'additional_director_email_id': False, 'deputy_director_email_id': False, 'zonal_head_email_id':False, 'file':False, 'priority':False , 'Upload Supporitng Documents':False})))
    total = total + len(allKeys_df)
    allKeys_df_dummy = pd.DataFrame(list(collection1.find({"Form": "Share inputs by your formation"}, {'additional_director_email_id': False,'deputy_director_email_id': False,'Estimated Value(Value in Crores-Rs)':False,'Suspect name':False,'gist_of_input': False, 'zonal_head_email_id':False, 'file':False, 'priority':False })))
    for doc in allKeys:
        pass
    headings=[]
    for val in list(doc.keys()):
        headings.append(str(val).capitalize())
    if doc:
        headings.sort()
        temp = headings[0]
        headings[0] = headings[1]
        headings[1] = temp    
        # print(headings)
        headings.remove("Regional_unit")
        headings.remove("Sub_regional_unit")
    allKeys_df_dummy_list = allKeys_df_dummy.to_dict(orient="records")
    page = request.args.get('page', 1)
    limit = 10
    # ROWS_PER_PAGE = 4
    offset = (int(page)) * limit-limit
    pagination = Pagination(page=page, per_page=limit, total=total, css_framework ="bootstrap5", offset=offset)
    dri_db = cluster["DRI"]
    if "formations" in list(dri_db.list_collection_names()):
        collection = dri_db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        zonal_units=formation_data.values.tolist()
    else:
        zonal_units=[]
    print(list(allKeys_df_dummy))
    return render_template("view_intelligence2.html",content=navbar_conditions(),zonal_units=zonal_units, filter_place=place_data, headings=headings,
                           pagination=pagination, allkeysdata=allKeys_df_dummy_list,allkeydataheadings = list(allKeys_df_dummy))

@app.route("/si_form")
def si_form():
    print("I am here in si_form")
    total_data = ["Cigarettes","Drugs","NDPS","Gold","Money"]
    dri_db = cluster["DRI"]
    if "formations" in list(dri_db.list_collection_names()):
        collection = dri_db["formations"]
        formation_data = pd.DataFrame(list(collection.find()))
        formation_data_list=formation_data.values.tolist()
    else:
        formation_data_list=[]
    regional_data = ['DRI Jaipur Regional Unit', 'DRI Regional Unit Noida', 'DRI Regional Unit Patna', 'DRI Tuticorin Regional Unit', 'DRI Gandhidham Regional Unit', 'DRI Surat Regional Unit', 'DRI Guwahati Regional Unit', 'DRI Siliguri Regional Unit', 'DRI, Silchar Regional Unit', 'DRI, Silchar Regional Unit', 'DRI Jammu Regional Unit', 'DRI Jammu Regional Unit', 'DRI Muzaffarpur Regional Unit', 'DRI Trichy Regional Unit', 'DRI Vijayawada Regional Unit', 'DRI Visakhapatnam Regional Unit', 'DRI Bhubaneswar Regional Unit', 'DRI Jamnagar Regional Unit', 'DRI Dimapur Regional Unit', 'DRI Agartala Regional Unit', 'DRI Agartala Regional Unit', 'DRI Aizwal Regional Unit', 'DRI Shillong Regional Unit', 'DRI Chandigarh Regional Unit', 'DRI Kannur Regional Unit', 'DRI Raipur Regional Unit', 'DRI Calicut Regional Unit', 'DRI Bhopal Regional Unit', 'DRI Mangalore Regional Unit', 'DRI Mangalore Regional Unit', 'DRI Madurai Sub-Regional Unit', 'DRI Siliguri Regional Unit']
    subregional_data =['Bhavnagar Sub Regional Unit', 'Vapi Sub Regional Unit', 'Madurai Sub Regional Unit', 'Jodhpur Sub Regional Unit', 'Nellore Sub Regional Unit', 'Kakinada Sub Regional Unit', 'Silchar Sub Regional Unit', 'Malda Sub Regional Unit', 'Dibrugarh Sub Regional Unit', 'Varanasi Sub Regional Unit', 'Purnea Sub Regional Unit', 'Karwar Sub Regional Unit', 'Belgaum Sub Regional Unit', 'Haldia Sub Regional Unit', 'Haldia Sub Regional Unit', 'Vapi Sub Regional Unit', 'Vapi Sub Regional Unit', 'Jodhpur Sub Regional Unit', 'Nellore Sub Regional Unit', 'Kakinada Sub Regional Unit', 'Silchar Sub Regional Unit', 'Silchar Sub Regional Unit', 'Dibrugarh Sub Regional Unit', 'Dibrugarh Sub Regional Unit', 'Varanasi Sub Regional Unit', 'Purnea Sub Regional Unit', 'Sri Nagar Sub Regional Unit', 'Karwar Sub Regional Unit', 'Belgaum Sub Regional Unit', 'Haldia Sub Regional Unit']
    return render_template("shareintelligence2.html",content=navbar_conditions(),buttons= button_conditions(),filter_commodity=total_data,filter_formation=formation_data_list,regional_data=regional_data,subregional_data=subregional_data)

@app.route("/edit_user",methods=["POST","GET"])
def edit_user():
    db = cluster["Authentication"]
    total_data_final = pd.DataFrame()
    message = None
    if "edit_success" in session:
        message = session.pop("edit_success",None)
    if "add_user_msg" in session:
        message = session.pop("add_user_msg",None)
    if "user_del_msg" in session:
        message = session.pop("user_del_msg",None)
    for col in list(db.list_collection_names()):
        if str(col)!="Admin":
            collection = db[str(col)]
            total_data = pd.DataFrame(list(collection.find()))
            total_data_final = pd.concat([total_data_final, total_data], ignore_index=True)
    # print(total_data_final)
    # user_data = pd.DataFrame(list(collection.find()))
    print(total_data_final.values.tolist())
    return render_template("users.html",content=navbar_conditions(), buttons=button_conditions(),designation=session["designation"],user_data=total_data_final.values.tolist(),message=message)

@app.route("/add_user")
def add_user():
    db = cluster["DRI"]
    collection = db["designation"]
    total_data = pd.DataFrame(list(collection.find()))
    collection = db["formations"]
    formation_data = pd.DataFrame(list(collection.find()))
    return render_template("profile.html",designation=session["designation"],content=navbar_conditions(), buttons=button_conditions(), filter_place = total_data.values.tolist(), filter_formation=formation_data.values.tolist())

@app.route("/subregional_unit")
def subregional_unit():
    db = cluster["DRI"]
    collection = db["subregional_unit"]
    message = None
    if "subregional_msg" in session:
        message = session.pop("subregional_msg",None)
    if "sub_regionalunits_del_msg" in session:
        message = session.pop("sub_regionalunits_del_msg",None)
    if "sub_reg_edit" in session:
        message = session.pop("sub_reg_edit",None)
    total_data = pd.DataFrame(list(collection.find()))
    return render_template("subregional_units.html",designation=session["designation"],content=navbar_conditions(), buttons=button_conditions(), total_data= total_data.to_dict(orient="records"),msg=message)

@app.route("/formations")
def formations():
    db = cluster["DRI"]
    collection = db["formations"]
    total_data = pd.DataFrame(list(collection.find()))
    message = None
    if "formations_msg" in session:
        message = session.pop("formations_msg",None)
    if "formations_edit_msg" in session:
        message = session.pop("formations_edit_msg",None)
    if "formations_del_msg" in session:
        message = session.pop("formations_del_msg",None)
    return render_template("formations.html",designation=session["designation"],content=navbar_conditions(), buttons=button_conditions(), total_data= total_data.values.tolist(),message=message)

@app.route("/commodities")
def commodities():
    db = cluster["DRI"]
    collection = db["commodity"]
    message = None
    if "commodities_del_msg" in session:
        message = session.pop("commodities_del_msg",None)
    total_data = pd.DataFrame(list(collection.find()))
    if "new_commodity_add_msg" in session:
        message = session.pop("new_commodity_add_msg",None)
    if "edit_msg" in session:
        message = session.pop("edit_msg",None)
    return render_template("commodity.html",designation=session["designation"],content=navbar_conditions(), buttons=button_conditions(), total_data= total_data.to_dict(orient="records"),msg=message)


@app.route("/designation")
def designation():
    db = cluster["DRI"]
    collection = db["designation"]
    total_data = pd.DataFrame(list(collection.find()))
    message =None
    if "new_db_added" in session:
        message = session.pop("new_db_added",None)
    if "designation_changed" in session:
        message = session.pop("designation_changed",None)
    return render_template("designation.html",designation=session["designation"], content=navbar_conditions(), buttons=button_conditions(),total_data= total_data.to_dict(orient="records"),msg=message)


@app.route("/DRI1_conditions", methods=["POST"])
def DRI1_conditions():
    try:
        button_input = request.form["submit_button"]
        if button_input == "DRI1":
            return redirect(url_for("DRI1"))
        elif button_input == "Home":
            return redirect(url_for("home"))
        elif button_input == "Formations":
            return redirect(url_for("formations"))
        elif button_input == "Commodity":
            return redirect(url_for("commodities"))
        elif button_input == "Designation":
            return redirect(url_for("designation"))
        elif button_input == "Regional Units":
            return redirect(url_for("regional_unit"))
        elif button_input == "Sub Regional Units":
            return redirect(url_for("subregional_unit"))
        elif button_input == "Share Inputs":
            return redirect(url_for("shareintelligence"))
        elif button_input == "Share Inputs Form":
            return redirect(url_for("si_form"))
        elif button_input == "DRI2":
            return redirect(url_for("DRI2"))
        elif button_input == "DRI3":
            return redirect(url_for("DRI3"))
        elif button_input == "User":
            return redirect(url_for("user_page"))
    except Exception as e:
        msg = f"invalid designation {e}"
        return render_template("home.html",msg=msg)
    
@app.route("/post_profile_settings", methods=["GET","POST"])
def post_profile_settings():
    print("post_profile_settings")
    if request.method == 'POST':
        person_name = request.form["person_name"]
        ssoid = request.form["ssoid"]
        email = request.form["email"]
        mobile_number = request.form["mobile_number"]
        designation = request.form["designation"]
        formation = request.form["formation"]
        create_password = request.form["create_password"]
        confirm_password = request.form["confirm_password"]
        if create_password == confirm_password:
            password = confirm_password
            db = cluster["Authentication"]
            collection = db[str(formation)]
            collection.insert_one({"ssoid": ssoid, "password": password, "email": email,"mobile number": mobile_number, "designation": designation,"formation":formation,"person_name":person_name})
        else:
            msg = f"invalid Details/Password"
            return render_template("home.html",msg=msg)
    session["add_user_msg"] = "Profile Saved"
    return redirect(url_for("edit_user"))

# @app.route("/commodity")
# def commodity():
#     return render_template("commodity.html",content=navbar_conditions(), buttons=button_conditions())
# @app.route("/designation")
# def designation():
#     return render_template("designation.html",content=navbar_conditions(), buttons=button_conditions())


def create_new_db(name):
    databases = list(cluster.list_database_names())
    for db in databases:
        if str(db) not in ["admin","local","config","DRI"]:
            database = cluster[str(db)]
            collection = database[str(name)]

@app.route("/formations_form", methods=["POST"])
def formations_form():
    db = cluster["DRI"]
    collection = db["formations"]
    if request.method == "POST":
        name = request.form["name"]
        address = request.form["address"]
        formation_form = {"Form":"formations_form","formation":name,"Address":address}
        create_new_db(name)
        collection.insert_one(formation_form)
        session["formations_msg"] = "Your form has been submitted"
        return redirect(url_for("formations"))
    
@app.route("/commodity_form", methods=["POST"])
def commodity_form():
    collection = db["commodity"]
    if request.method == "POST":
        commodity = request.form["commodity"]
        commodity_form = {"Form":"commodity_form","commodity":commodity}
        collection.insert_one(commodity_form)
        session["new_commodity_add_msg"] = "Your form has been submitted"
    return redirect(url_for("commodities"))
    


@app.route("/designation_form", methods=["POST"])
def designation_form():
    db = cluster["DRI"]
    collection = db["designation"]
    if request.method == "POST":
        designation = request.form["designation"]
        designation_form = {"Form":"designation_form","designation":designation}
        collection.insert_one(designation_form)
        session["new_db_added"] = "Your form has been submitted"
    return redirect(url_for("designation"))    

@app.route('/delete', methods=['POST'])
def delete():
    if request.method == "POST":
        id_value = request.form.get("id_value")
        id_value = ObjectId(id_value)
        collection_id = request.form.get("collection_id")
        print(collection_id)
        try:
            if collection_id in ["commodity","designation","regional_unit","subregional_unit","formations"]:
                db = cluster["DRI"]
                input_data = db[collection_id]
                if collection_id == "commodity":
                    if input_data.find({"_id":id_value}):
                        input_data.delete_one({'_id': id_value})
                    session["commodities_del_msg"] = "Successfully Commodity Removed"
                    return redirect(url_for("commodities"))
                elif collection_id == "designation":
                    if input_data.find({"_id":id_value}):
                        input_data.delete_one({'_id': id_value})
                    session["designation_del_msg"] = "Successfully Designation Removed"
                    return redirect(url_for("designation"))
                elif collection_id == "regional_unit":
                    if input_data.find({"_id":id_value}):
                        input_data.delete_one({'_id': id_value})
                    session["regionalunits_del_msg"] = "Successfully Regional unit Removed"
                    return redirect(url_for("regional_unit"))
                elif collection_id == "subregional_unit":
                    if input_data.find({"_id":id_value}):
                        input_data.delete_one({'_id': id_value})
                    session["sub_regionalunits_del_msg"] = "Successfully Sub Regional unit Removed"
                    return redirect(url_for("subregional_unit"))
                elif collection_id == "formations":
                    if input_data.find({"_id":id_value}):
                        input_data.delete_one({'_id': id_value})
                    session["formations_del_msg"] = "Successfully Formations Removed"
                    return redirect(url_for("formations"))
            elif collection_id == "authentication":
                user_collection = request.form.get("collection_id_main")
                print(user_collection)
                database_list = list(cluster.list_database_names())
                for db in database_list:
                    database = cluster[str(db)]
                    input_data = database[str(user_collection)]
                    if input_data.find_one({"_id":id_value}):
                        input_data.delete_one({"_id":id_value})
                session["user_del_msg"] = "Successfully Removed The User"
                return redirect(url_for("edit_user"))

             


        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        

@app.route('/edit', methods=['POST'])
def edit():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        new_designation = request.form["new_designation"]
        id_value = ObjectId(id_value)
        collection_id = request.form.get("collection_id")
        input_data = db[collection_id]
        try:
            if input_data.find({"_id":id_value}):
                input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"designation":new_designation}})
                session["designation_changed"] = "Designation changed successfully"
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        return redirect(url_for("designation"))
        

@app.route('/edit_commodity', methods=['POST'])
def edit_commodity():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        new_commodity = request.form["new_commodity"]
        id_value = ObjectId(id_value)
        print(id_value)
        collection_id = request.form.get("collection_id")
        print(collection_id)
        input_data = db[collection_id]
        try:
            if input_data.find({"_id":id_value}):
                input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"commodity":new_commodity}})
                session["edit_msg"] = "commodity changed successfully"
            return redirect(url_for("commodities"))
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        
@app.route('/edit_regionalunit', methods=['POST'])
def edit_regionalunit():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        new_regional_unit = request.form["new_regional_unit"]
        id_value = ObjectId(id_value)
        print(id_value)
        collection_id = request.form.get("collection_id")
        print(collection_id)
        input_data = db[collection_id]
        try:
            if input_data.find({"_id":id_value}):
                input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"regional_unit":new_regional_unit}})
                session["r_u_edit_msg"] = "regional unit changed successfully"
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        return redirect(url_for("regional_unit"))
        

@app.route('/edit_subregionalunit', methods=['POST'])
def edit_subregionalunit():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        new_subregional_unit = request.form["new_subregional_unit"]

        id_value = ObjectId(id_value)
        print(id_value)
        collection_id = request.form.get("collection_id")
        print(collection_id)
        input_data = db[collection_id]
        try:
            if input_data.find({"_id":id_value}):
                input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"subregional_unit":new_subregional_unit}})
                session["sub_reg_edit"] = "Sub regional unit changed successfully"
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        return redirect(url_for("subregional_unit"))
        
def db_name_replace(old_name,new_name):
    db_list = list(cluster.list_database_names())
    for db in db_list:
        if str(db) not in ["admin","local","config","DRI"]:
            database = cluster[str(db)]
            for col in list(database.list_collection_names()):
                if col==old_name and col!=new_name:
                    collection = database[str(col)]
                    collection.update_many({},{ "$set": { "formation": new_name } })
                    collection.rename(new_name, dropTarget = True)

@app.route('/edit_formation', methods=['POST'])
def edit_formation():
    db = cluster["DRI"]
    if request.method == "POST":
        id_value = request.form["id_value"]
        new_name = request.form["new_name"]
        new_address = request.form["new_address"]
        old_name = request.form["old_col_name"]
        id_value = ObjectId(id_value)
        collection_id = request.form.get("collection_id")
        input_data = db[collection_id]
        try:
            if input_data.find({"_id":id_value}):
                if len(str(new_name))!=0 and len(str(new_address))==0:
                    db_name_replace(old_name,new_name)
                    input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"formation":new_name, "Address":new_address}})
                elif len(str(new_address))!=0:
                    if len(str(new_name))==0:
                        input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"Address":new_address}})
                    elif len(str(new_name))!=0:
                        db_name_replace(old_name,new_name)
                        input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"formation":new_name,"Address":new_address}})
                session["formations_edit_msg"] = "formation changed successfully"
            return redirect(url_for("formations"))
        except Exception as e:
            print(e)
            return render_template("home.html", msg=f"{e}")
        
@app.route('/users', methods=['POST'])
def users():
    collection = db["authentication"]
    user_data = pd.DataFrame(list(collection.find()))
    return render_template("users.html",content=navbar_conditions(),buttons=button_conditions(),user_data=user_data.values.tolist())

@app.route('/edit_user_details', methods=['POST'])
def edit_user_details():
    db = cluster["Authentication"]
    if request.method == "POST":
        id_value = request.form.get("id_value")
        new_password = request.form["new_password"]
        new_formation = request.form["new_formation"]
        collection_id = request.form.get("collection_id") 
        id_value = ObjectId(id_value)
        input_data = db[str(collection_id)]          
        try:
            if input_data.find({"_id":id_value}):
                print("yesss")
                if len(str(new_formation))==0 and len(str(new_password))!=0:
                    input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"password":new_password}})
                elif len(str(new_formation))!=0:
                    new_input_data = db[str(new_formation)]
                    docs=input_data.find({"_id":ObjectId(id_value)})
                    for doc in docs:
                        new_input_data.insert_one(doc)
                        input_data.delete_one(doc)
                    if len(str(new_password))==0:
                        new_input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"formation":new_formation}})
                    elif len(str(new_password))!=0:
                        new_input_data.update_one({"_id":ObjectId(id_value)},{ "$set": {"password":new_password,"formation":new_formation}})
                session["edit_success"] = "Data changed successfully"
            user_data = pd.DataFrame(list(input_data.find()))
            user_data_for_col = pd.DataFrame()
            user_data_for_col = pd.concat([user_data_for_col, user_data], ignore_index=True)
            for col in list(db.list_collection_names()):
                if str(col)!=collection_id:
                    col= db[str(col)]
                    user_data = pd.DataFrame(list(col.find()))
                    user_data_for_col = pd.concat([user_data_for_col, user_data], ignore_index=True)
            return redirect(url_for("edit_user"))
        except Exception as e:
            print(e)
            return render_template("home.html", content=navbar_conditions(),buttons=button_conditions(),msg=f"{e}")
        

@app.route("/regional_unit")
def regional_unit():
    collection = db["regional_unit"]
    total_data = pd.DataFrame(list(collection.find()))
    message = None
    if "regional_msg" in session:
        message = session.pop("regional_msg",None)
    if "regionalunits_del_msg" in session:
        message = session.pop("regionalunits_del_msg",None)
    if "r_u_edit_msg" in session:
        message = session.pop("r_u_edit_msg",None)
    return render_template("regionalunits.html",content=navbar_conditions(),buttons=button_conditions(),designation=session["designation"], total_data= total_data.to_dict(orient="records"),msg = message)

@app.route("/regional_unit_display", methods=["POST","GET"])
def regional_unit_display():
    collection = db["regional_unit"]
    if request.method == "POST":
        regional_unit = request.form["regional_unit"]
        regional_unit = {"Form":"regional_unit","regional_unit":regional_unit}
        collection.insert_one(regional_unit)
    session["regional_msg"] = "Your form has been submitted"
    return redirect(url_for("regional_unit"))

@app.route("/subregional_unit_display", methods=["POST"])
def subregional_unit_display():
    db = cluster["DRI"]
    collection = db["subregional_unit"]
    if request.method == "POST":
        subregional_unit = request.form["subregional_unit"]
        subregional_unit = {"Form":"subregional_unit","subregional_unit":subregional_unit}
        collection.insert_one(subregional_unit)
        session["subregional_msg"] = "Your form has been submitted"
        return redirect(url_for("subregional_unit"))
        
if __name__ == "__main__":
    app.run(debug=True, )
